#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成：中国大学生计算机设计大赛 作品信息概要表（大数据应用，2026版）
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"d:\桌面\xuezhi-portrait-master\作品信息概要表_大数据应用_2026版.docx"


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '<w:top w:val="{top_val}" w:sz="{top_sz}" w:space="0" w:color="{color}"/>'
        '<w:left w:val="{left_val}" w:sz="{left_sz}" w:space="0" w:color="{color}"/>'
        '<w:bottom w:val="{bottom_val}" w:sz="{bottom_sz}" w:space="0" w:color="{color}"/>'
        '<w:right w:val="{right_val}" w:sz="{right_sz}" w:space="0" w:color="{color}"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    # 替换占位符
    top_val = kwargs.get('top', {}).get('val', 'single')
    top_sz = kwargs.get('top', {}).get('sz', '4')
    left_val = kwargs.get('left', {}).get('val', 'single')
    left_sz = kwargs.get('left', {}).get('sz', '4')
    bottom_val = kwargs.get('bottom', {}).get('val', 'single')
    bottom_sz = kwargs.get('bottom', {}).get('sz', '4')
    right_val = kwargs.get('right', {}).get('val', 'single')
    right_sz = kwargs.get('right', {}).get('sz', '4')
    color = kwargs.get('color', '000000')

    xml_str = (
        '<w:tcBorders %s>'
        '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '</w:tcBorders>' % (
            nsdecls('w'),
            top_val, top_sz, color,
            left_val, left_sz, color,
            bottom_val, bottom_sz, color,
            right_val, right_sz, color
        )
    )
    tcPr.append(parse_xml(xml_str))


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(
        '<w:shd %s w:fill="%s" w:val="clear"/>' % (nsdecls('w'), color)
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_formatted_paragraph(doc, text, bold=False, size=11, font_name='宋体',
                             alignment=None, space_after=6, space_before=0,
                             first_line_indent=None, color=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)

    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_table_row(table, cells_data, bold=False, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, (text, width_cm) in enumerate(cells_data):
        cell = row.cells[i]
        cell.width = Cm(width_cm)
        # 清除默认段落
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = bold
        if header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'D9E2F3')
        # 设置边距
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            '<w:tcMar %s>'
            '<w:top w:w="40" w:type="dxa"/>'
            '<w:left w:w="80" w:type="dxa"/>'
            '<w:bottom w:w="40" w:type="dxa"/>'
            '<w:right w:w="80" w:type="dxa"/>'
            '</w:tcMar>' % nsdecls('w')
        )
        tcPr.append(tcMar)
    return row


def create_document():
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 标题 ====================
    title = add_formatted_paragraph(
        doc,
        '中国大学生计算机设计大赛（2026版）\n作品信息概要表',
        bold=True, size=16, font_name='黑体',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    subtitle = add_formatted_paragraph(
        doc,
        '（大数据应用赛道）',
        bold=False, size=12, font_name='宋体',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # ==================== 基本信息表 ====================
    basic_info = [
        ('作品名称', '学智画像——高校计算机能力画像与智能学习资源生成系统'),
        ('参赛类别', '大数据应用'),
        ('参赛学校', '【请填写学校全称】'),
        ('作品类型', 'Web应用（Flask + 多智能体协同 + 知识库RAG）'),
    ]

    table1 = doc.add_table(rows=0, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = True

    for label, value in basic_info:
        row = table1.add_row()
        # 标签列
        cell0 = row.cells[0]
        cell0.width = Cm(3.5)
        cell0.paragraphs[0].clear()
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(10.5)
        r0.font.name = '宋体'
        r0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell0, 'D9E2F3')

        # 值列
        cell1 = row.cells[1]
        cell1.width = Cm(12.5)
        cell1.paragraphs[0].clear()
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(10.5)
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()  # 空行

    # ==================== 团队成员 ====================
    add_formatted_paragraph(
        doc, '一、团队成员信息', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    team_table = doc.add_table(rows=1, cols=5)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['序号', '姓名', '专业/年级', '分工职责', '联系方式']
    header_row = team_table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    # 3行占位 (最多5人)
    for n in range(1, 6):
        row = team_table.add_row()
        for i, val in enumerate([str(n), '【姓名】', '【专业/年级】', '【分工】', '【联系方式】']):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 指导教师 ====================
    add_formatted_paragraph(
        doc, '二、指导教师信息', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    advisor_table = doc.add_table(rows=1, cols=4)
    advisor_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    advisor_headers = ['序号', '姓名', '所属单位', '联系方式']
    for i, h in enumerate(advisor_headers):
        cell = advisor_table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    for n in range(1, 3):
        row = advisor_table.add_row()
        for i, val in enumerate([str(n), '【姓名】', '【所属单位】', '【联系方式】']):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 作品简介 ====================
    add_formatted_paragraph(
        doc, '三、作品简介（500字以内）', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    intro_text = (
        '学智画像是一款面向高校计算机教育的智能学习辅助系统，基于Flask Web框架构建，集成多智能体协同架构与知识库RAG（检索增强生成）技术。'
        '系统服务于教师、学生和管理员三类角色，核心功能包括：\n\n'
        '（1）多维学习画像：通过10维度评估模型（编程能力、算法思维、数学基础、英语水平、计算机网络、民航特色知识等），'
        '结合KMeans聚类和PCA降维，为学生生成个性化学习画像；\n'
        '（2）知识库与RAG反幻觉：支持教师上传课程知识文档（Markdown/CSV/JSON/TXT），'
        '自动解析分块并建立TF-IDF检索索引，智能体基于检索结果生成带引用的问答，有效抑制大模型幻觉；\n'
        '（3）多智能体协同：7个专业智能体（画像构建、知识库管理、资源生成、学习规划、辅导、评估、协调）'
        '协同工作，实现从能力诊断到个性化学习的完整闭环；\n'
        '（4）智能资源生成：集成讯飞智文API、Coze工作流等，支持PPT课件、思维导图、练习题、数字人视频等'
        '多模态学习资源的自动生成；\n'
        '（5）大数据分析：利用RandomForest随机森林进行学习预警预测，时间序列分析追踪学习趋势，'
        '关联规则挖掘发现知识点间的内在联系。'
    )

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.first_line_indent = Cm(0.74)
    p_intro.paragraph_format.line_spacing = 1.5
    r = p_intro.add_run(intro_text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 核心技术 ====================
    add_formatted_paragraph(
        doc, '四、核心技术说明', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    tech_items = [
        ('后端框架：', 'Python Flask + SQLAlchemy ORM + Flask-Login认证'),
        ('前端技术：', 'Jinja2模板引擎 + Bootstrap 5响应式布局 + Chart.js/ECharts可视化'),
        ('机器学习/AI：', 'scikit-learn（KMeans聚类、RandomForest预测、PCA降维）；'
         'PyTorch深度学习（知识点掌握预测模型、学习风格分类器、注意力机制）；'
         '7智能体协同框架（画像构建/知识库管理/资源生成/学习规划/辅导/评估/协调）'),
        ('知识库RAG：', '文档自动解析分块 + TF-IDF向量检索 + 引用式反幻觉问答'),
        ('外部AI集成：', '讯飞智文PPT生成API、讯飞虚拟数字人API、Coze工作流API'),
        ('数据分析：', '大数据分析引擎（时间序列分析/聚类/预测/关联规则挖掘）'),
        ('部署方案：', 'Gunicorn/Waitress生产部署 + Docker容器化 + Nginx反向代理'),
    ]

    for label, detail in tech_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = '宋体'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2 = p.add_run(detail)
        r2.font.size = Pt(10.5)
        r2.font.name = '宋体'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 数据集说明 ====================
    add_formatted_paragraph(
        doc, '五、数据集说明及下载链接', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    dataset_intro = (
        '本项目采用了以下数据集，在提交的源码包中仅保留少量典型样本，完整数据集'
        '可通过以下链接下载：'
    )
    p_ds = doc.add_paragraph()
    p_ds.paragraph_format.first_line_indent = Cm(0.74)
    p_ds.paragraph_format.line_spacing = 1.5
    r = p_ds.add_run(dataset_intro)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 数据集表
    ds_table = doc.add_table(rows=1, cols=5)
    ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    ds_headers = ['数据集名称', '数据说明', '原始大小', '包内样本', '完整下载链接']
    for i, h in enumerate(ds_headers):
        cell = ds_table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')

    datasets = [
        ['quiz_data.csv\n（模拟答题数据）',
         '50名模拟学生的答题记录，包含题目ID、答案、用时、得分等字段',
         '520 KB\n（~3000条记录）',
         '50条样本\n(quiz_data_sample.csv)',
         '【请上传至网盘并填写链接】\n推荐：百度网盘/阿里云盘'],
        ['计算机组成原理\n知识库样例包',
         '课程导学、核心知识讲义、实验指导、章节题库、术语表、复习提纲等6类文档',
         '10 KB\n（7个文件）',
         '完整保留\n(data/knowledge_base/)',
         '（已完整包含在源码包中）'],
    ]

    for ds_row in datasets:
        row = ds_table.add_row()
        for i, val in enumerate(ds_row):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    note_text = (
        '注：quiz_data.csv为基于真实考试场景模拟生成的数据集，用于系统功能演示与测试。'
        '若评审需要，可提供数据生成Python脚本（见源码包 app/utils/import_data.py 和 set_db.py）。'
        '完整quiz_data.csv数据集（520KB）的下载链接请填入上方表格中。'
    )
    p_note = doc.add_paragraph()
    p_note.paragraph_format.first_line_indent = Cm(0.74)
    r = p_note.add_run(note_text)
    r.font.size = Pt(9)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ==================== 创新点 ====================
    add_formatted_paragraph(
        doc, '六、主要创新点', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    innovations = [
        '对话式学习画像构建：突破传统问卷模式，通过智能体对话交互获取学生信息，结合10维度评估模型自动生成个性化画像，用户体验更自然。',
        '知识库RAG反幻觉机制：针对大模型"幻觉"问题，设计文档解析-分块-检索-引用的完整pipeline，所有AI回答均附带原文引用来源，有效提升可信度。',
        '多智能体协同架构：7个专业化智能体各司其职又协同配合，通过协调器（Coordinator）统一调度，形成完整的"诊断-规划-学习-评估"闭环。',
        '多模态资源自动生成：整合讯飞智文、Coze等外部AI能力，一键生成PPT课件、思维导图、数字人讲解视频等6类学习资源，大幅降低教师备课负担。',
        '大数据驱动的预警与推荐：基于RandomForest预测高风险学生，结合时间序列分析追踪学习趋势，实现从"事后补救"到"事前预防"的转变。',
    ]

    for i, item in enumerate(innovations, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        r = p.add_run(f'{i}. {item}')
        r.font.size = Pt(10.5)
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()

    # ==================== 部署说明 ====================
    add_formatted_paragraph(
        doc, '七、系统部署与运行说明', bold=True, size=12, font_name='黑体',
        space_after=6
    )

    deploy_text = (
        'Python 3.8+环境，安装依赖：pip install -r requirements.txt\n'
        '初始化数据库并生成测试数据：python set_db.py\n'
        '启动开发服务器：python run.py\n'
        '访问地址：http://localhost:5000\n'
        '默认测试账号：admin@test.com（管理员）/ teacher@test.com（教师）/ student@test.com（学生），密码均为123456\n\n'
        '详细部署说明请参考源码包中的 README.md 和 QUICK_START.md。'
    )
    p_deploy = doc.add_paragraph()
    p_deploy.paragraph_format.first_line_indent = Cm(0.74)
    p_deploy.paragraph_format.line_spacing = 1.5
    r = p_deploy.add_run(deploy_text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 页脚声明 ====================
    doc.add_paragraph()
    doc.add_paragraph()
    footer_text = (
        '声明：本作品所有源代码均为团队自主开发。作品中引用的开源组件（Flask、scikit-learn、'
        'PyTorch、Bootstrap、Chart.js等）均已在源码包中 docs/A3_open_source_notice.md 文件中'
        '详细列出其名称、版本、许可证类型及用途说明。'
    )
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(20)
    r = p_footer.add_run(footer_text)
    r.font.size = Pt(9)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.color.rgb = RGBColor(100, 100, 100)

    # ==================== 落款 ====================
    sign_text = '填表日期：2026年____月____日'
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sign.add_run(sign_text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存
    doc.save(OUTPUT_PATH)
    print(f'[OK] 作品信息概要表已生成: {OUTPUT_PATH}')
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f'     文件大小: {size_kb:.1f} KB')


if __name__ == '__main__':
    create_document()
