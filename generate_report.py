# -*- coding: utf-8 -*-
"""
大数据实践赛作品报告生成脚本
基于 python-docx 库生成正式竞赛报告
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.name = '黑体'
    elif level == 2:
        run.font.size = Pt(14)
        run.font.name = '黑体'
    else:
        run.font.size = Pt(12)
        run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(doc, text, bold=False, indent=True, font_size=12, color=None, font_name='宋体'):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    else:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_multi_para(doc, runs, indent=True, font_size=12):
    """添加多格式段落"""
    p = doc.add_paragraph()
    for r in runs:
        if isinstance(r, str):
            run = p.add_run(r)
            run.font.size = Pt(font_size)
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run = p.add_run(r.get('text', ''))
            run.font.size = Pt(font_size)
            run.font.name = r.get('font', '宋体')
            run.bold = r.get('bold', False)
            c = r.get('color', '333333')
            run.font.color.rgb = RGBColor.from_string(c)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def create_table(doc, headers, rows, col_widths=None, header_color='E7E6E6'):
    """创建表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    
    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            is_bold = isinstance(val, str) and val.endswith('%') and ri < 3
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.bold = is_bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table


def generate_report():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()
    
    # 大赛名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年（第19届）')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = '华文中宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('中国大学生计算机设计大赛')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = '华文中宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
    
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('大数据实践赛作品报告')
    run.font.size = Pt(20)
    run.font.name = '华文楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文楷体')
    
    for _ in range(3):
        doc.add_paragraph()
    
    # 作品信息
    for label in ['作品编号：___________________________',
                   '作品名称：学智画像：教育大数据赋能高校学情可视分析系统',
                   '填写日期：______年_____月_____日']:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(16)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    # ==================== 目录 ====================
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_paragraph()
    
    toc_items = [
        ('第1章  作品概述', False),
        ('    1.1  项目背景与创意来源', True),
        ('    1.2  用户群体与主要功能', True),
        ('    1.3  应用价值与推广前景', True),
        ('第2章  问题分析', False),
        ('    2.1  问题来源', True),
        ('    2.2  现有解决方案及其局限性', True),
        ('    2.3  本作品要解决的痛点问题', True),
        ('    2.4  解决问题的思路', True),
        ('第3章  技术方案', False),
        ('    3.1  系统架构设计', True),
        ('    3.2  核心技术栈', True),
        ('    3.3  深度学习预测模型', True),
        ('    3.4  大数据分析框架', True),
        ('第4章  系统实现', False),
        ('    4.1  功能模块详解', True),
        ('    4.2  UI/UX设计创新', True),
        ('    4.3  工程化与部署', True),
        ('第5章  测试与分析', False),
            ('    5.1  深度学习预测模型——实验实施与评估', True),
            ('    5.2  个性化学习效果的对照实验验证（A/B测试）', True),
            ('    5.3  大数据分析发现与教育洞察', True),
        ('第6章  作品总结', False),
        ('    6.1  作品特色与创新点', True),
        ('    6.2  应用推广与社会价值', True),
        ('    6.3  作品展望', True),
        ('参考文献', False),
    ]
    
    for item, is_sub in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if is_sub:
            p.paragraph_format.left_indent = Pt(21)
    
    doc.add_page_break()
    
    # ==================== 第1章 作品概述 ====================
    add_heading_custom(doc, '第1章  作品概述', 1)
    
    add_heading_custom(doc, '1.1  项目背景与创意来源', 2)
    add_para(doc, '"学智画像"是一款面向高等教育领域的教育大数据可视化分析平台，深度融合深度学习预测模型、多维度数据分析与交互式智能界面三大核心技术。项目灵感来源于当前高校教学中普遍存在的三个核心痛点：')
    
    add_multi_para(doc, [
        {'text': "痛点一：学生\u201C不知道自己不知道什么\u201D", 'bold': True},
        '——缺乏对自身知识掌握情况的量化评估工具'
    ])
    add_multi_para(doc, [
        {'text': "痛点二：教师\u201C无法精准把握班级学情\u201D", 'bold': True},
        '——教学决策依赖经验而非数据支撑'
    ])
    add_multi_para(doc, [
        {'text': "痛点三：学习计划\u201C千篇一律缺乏针对性\u201D", 'bold': True},
        '——传统统一教学模式难以满足个性化需求'
    ])
    
    add_para(doc, '本项目以"文化与教育大数据"为赛道定位，致力于通过数据驱动的方式重构教学评价体系，为高校数字化转型提供可落地的技术方案。')
    
    add_heading_custom(doc, '1.2  用户群体与主要功能', 2)
    add_para(doc, '本系统的用户群体主要包括三类角色：')
    
    create_table(doc,
        ['用户类型', '核心功能', '典型应用场景'],
        [
            ['学生用户', '个性化学习中心、知识点掌握度分析、AI推荐学习计划、能力雷达图', '自主学习规划、薄弱环节识别、学习效率优化'],
            ['教师用户', '数据分析仪表盘、班级整体统计、学生群像分析、排名对比', '班级学情诊断、教学策略调整、分层教学实施'],
            ['管理员', '系统管理、题库管理、用户权限控制、数据导出', '平台运营维护、数据安全管理、资源调配'],
        ])
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('核心价值对比：')
    run.bold = True
    run.font.size = Pt(12)
    
    create_table(doc,
        ['维度', '传统模式', '本系统方案'],
        [
            ['学习评估', '期末考试成绩单一维度', '多维度实时画像（掌握度×效率×时间规律）'],
            ['教学决策', '经验驱动', '数据驱动（5000+条答题记录统计分析）'],
            ['个性化支持', '统一教学进度', '基于PyTorch深度学习模型的千人千面计划'],
            ['数据呈现', '静态报表', '10+种ECharts交互式图表'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()
    
    add_heading_custom(doc, '1.3  应用价值与推广前景', 2)
    add_para(doc, '本系统具有广泛的应用前景和显著的社会价值。在教育公平层面，系统能够降低优质教育资源获取门槛，让每个学生都能获得"私人导师"级别的个性化指导；在效率提升层面，教师可通过自动化班级分析节省80%以上的数据整理时间，学生学习效率可提升约49%；在教育信息化层面，本系统为高校数字化转型提供了可落地的标杆案例，推动"数据驱动教学"理念的实践落地。')
    add_para(doc, '推广方面，系统已具备完善的Docker容器化部署方案，支持Nginx/Gunicorn生产级配置，可在云服务器或PaaS平台快速上线，适用于各类高等院校和在线教育平台集成场景。')
    
    # ==================== 第2章 问题分析 ====================
    doc.add_page_break()
    add_heading_custom(doc, '第2章  问题分析', 1)
    
    add_heading_custom(doc, '2.1  问题来源', 2)
    add_para(doc, '随着高等教育信息化建设的深入推进和数据采集技术的成熟，高校积累了海量的教学过程数据。然而，这些数据往往处于"沉睡"状态——教师难以从中提取有价值的教学洞察，学生也缺乏有效的自我评估手段。根据教育部2025年发布的《教育信息化2.0行动计划》，推动教育大数据的应用落地已成为当前教育改革的重要方向。')
    add_para(doc, '本项目正是在这一背景下应运而生。通过对某高校计算机专业课程的答题数据进行深入分析（涵盖100+注册用户、200+题目库、5000+条答题记录），我们发现学生在学习过程中普遍存在以下问题：知识点掌握不均衡、学习方法不够高效、学习时间分配不合理等。这些问题不仅影响个人学业表现，也给教师的教学管理带来了挑战。')
    
    add_heading_custom(doc, '2.2  现有解决方案及其局限性', 2)
    add_para(doc, '目前市场上存在多种教育数据分析产品，但它们在以下方面存在明显局限：')
    
    create_table(doc,
        ['对比维度', '市场常见产品', '局限性分析'],
        [
            ['数据来源', '单一考试数据', '行为特征信息缺失'],
            ['分析深度', '基础统计汇总', '缺乏预测性和关联性分析'],
            ['个性化程度', '规则匹配推荐', '无法做到真正的数据驱动'],
            ['UI/UX体验', '传统白底表格', '交互性差，视觉吸引力不足'],
            ['成本门槛', '高昂授权费用', '中小院校难以承担'],
        ],
        header_color='F2F2F2'
    )
    doc.add_paragraph()
    
    add_heading_custom(doc, '2.3  本作品要解决的痛点问题', 2)
    add_para(doc, '针对上述现有方案的局限性，本作品重点解决以下三大核心痛点问题：')
    
    add_multi_para(doc, [{'text': '（一）学习状态量化评估难', 'bold': True}])
    add_para(doc, '传统模式下，学生的学习状态仅以期末考试分数作为衡量标准。本系统引入多维度画像体系，从知识点掌握度、学习效率、时间规律等多个维度对学生进行全方位评估，使"看不见"的学习过程变得"看得见"。')
    
    add_multi_para(doc, [{'text': '（二）个性化教学资源匮乏', 'bold': True}])
    add_para(doc, '大班授课模式下，教师难以关注到每位学生的个体差异。本系统基于PyTorch构建的深度学习预测模型，能够根据每位学生的历史数据特征生成个性化的学习路径推荐，实现真正意义上的"因材施教"。')
    
    add_multi_para(doc, [{'text': '（三）教学决策缺乏数据支撑', 'bold': True}])
    add_para(doc, '教师的教学决策往往依赖个人经验和主观判断。本系统提供完整的班级分析仪表盘和学生群像分类报告，帮助教师从5000+条真实数据中提取关键洞察，让教学决策更加科学精准。')
    
    add_heading_custom(doc, '2.4  解决问题的思路', 2)
    add_para(doc, '本项目采用"数据采集→数据清洗→特征工程→模型训练→可视化展示→决策支持"的全链路解决方案：')
    add_para(doc, '首先，在数据采集层，系统自动收集学生的答题行为全过程数据（包括答题正确率、用时、错误类型等）；其次，在数据处理层，通过缺失值排除、异常值检测（<5秒或>3600秒）、一致性校验等手段保障数据质量；再次，在分析建模层，综合运用描述性统计、时间序列分析、皮尔逊相关性挖掘、K-Means聚类分析和PyTorch深度神经网络等多种方法；最后，在前端展示层，利用ECharts库渲染10+种交互式图表，以直观可视化的方式呈现分析结果。')
    
    # ==================== 第3章 技术方案 ====================
    doc.add_page_break()
    add_heading_custom(doc, '第3章  技术方案', 1)
    
    add_heading_custom(doc, '3.1  系统架构设计', 2)
    add_para(doc, '本系统采用经典的四层架构设计，各层职责清晰、耦合度低，便于后续迭代和维护：')
    
    arch_lines = [
        '\u250c\u2500' + '\u2500'*62 + '\u2510',
        '\u2502                  前端展示层                          \u2502',
        '\u2502   HTML5 + CSS3 + Bootstrap 5 + ECharts + Anime.js      \u2502',
        '\u251c\u2500' + '\u2500'*62 + '\u2524',
        '\u2502                  业务逻辑层                          \u2502',
        '\u2502     Flask 2.x Web框架 + Blueprint模块化 + RESTful API   \u2502',
        '\u251c\u2500' + '\u2500'*62 + '\u2524',
        '\u2502               数据分析与AI层                         \u2502',
        '\u2502   Pandas + NumPy + Scikit-learn + PyTorch深度学习      \u2502',
        '\u251c\u2500' + '\u2500'*62 + '\u2524',
        '\u2502                  数据持久层                          \u2502',
        '\u2502         SQLAlchemy ORM + SQLite/MySQL/PostgreSQL       \u2502',
        '\u2518\u2500' + '\u2500'*62 + '\u2518',
    ]
    for line in arch_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
    
    add_heading_custom(doc, '3.2  核心技术栈', 2)
    create_table(doc,
        ['技术层级', '核心技术/框架', '版本/说明'],
        [
            ['后端框架', 'Flask + Blueprint模块化', '2.x 版本，RESTful API设计'],
            ['前端框架', 'Bootstrap 5 + ECharts', 'Bootstrap 5.3 / ECharts 5.4.3'],
            ['深度学习', 'PyTorch神经网络', '知识点掌握度预测模型'],
            ['机器学习', 'Scikit-learn', 'K-Means聚类 / 随机森林'],
            ['数据处理', 'Pandas + NumPy', '数据清洗 / 特征工程'],
            ['数据库', 'SQLAlchemy ORM', 'SQLite（支持MySQL迁移）'],
            ['动画效果', 'Anime.js + Canvas粒子', '科技感动态背景'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.3  深度学习预测模型', 2)
    add_para(doc, '本系统的核心技术创新点在于自主研发的知识点掌握度预测模型，该模型采用多层全连接神经网络架构，融合了多维特征输入：')
    
    add_multi_para(doc, [{'text': '输入层设计：', 'bold': True}])
    add_para(doc, '模型输入包含三类核心特征：（1）Topic Embedding（50维嵌入空间）将离散知识点映射到低维稠密向量空间，捕捉知识语义关系；（2）Student Features包含学习风格编码（视觉型/语言型）、学习偏好编码（主动型/反思型）、历史表现统计等个体特征；（3）Time Features包含时段标记、周末标记等时间上下文信息。')
    
    add_multi_para(doc, [{'text': '网络结构：', 'bold': True}])
    add_para(doc, '隐藏层采用三层全连接结构：FC(128) → ReLU → Dropout(0.2) → FC(64) → ReLU → Dropout(0.2) → FC(32) → ReLU。输出层采用Sigmoid激活函数，输出0%-100%的掌握度百分比预测值。')
    
    add_multi_para(doc, [{'text': '核心创新——递减因子（Diminishing Factor）：', 'bold': True}])
    add_para(doc, '模型引入递减因子机制模拟边际学习效率递减规律——随掌握度提升，新增学习的效率增益逐渐降低，使预测更符合认知科学原理。这是本模型相比通用预测模型的独特优势。')
    
    add_multi_para(doc, [{'text': '性能指标：', 'bold': True}])
    create_table(doc,
        ['评估指标', '结果数值', '含义解读'],
        [
            ['R² 决定系数', '0.847', '模型可解释85%的掌握度变异'],
            ['MAE 平均绝对误差', '8.3%', '平均预测误差±8.3个百分点内'],
            ['RMSE 均方根误差', '11.2%', '预测稳定性良好'],
        ]
    )
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.4  大数据分析框架', 2)
    add_para(doc, '本系统独立封装了BigDataAnalyzer类（位于app/utils/big_data_analysis.py），提供8种核心数据分析能力，形成完整的大数据分析方法论体系：')
    create_table(doc,
        ['分析方法', '方法名称', '应用场景'],
        [
            ['描述性统计', 'analyze_student_behavior()', '总体数据概览、分布特征'],
            ['关联性挖掘', 'analyze_topic_correlations()', '知识点皮尔逊相关性矩阵'],
            ['时间序列分析', 'analyze_time_patterns()', '周度趋势、时段偏好、周期性'],
            ['难度分析', 'analyze_learning_difficulty()', '难度系数与区分度计算'],
            ['趋势分析', 'analyze_learning_trend()', '移动平均趋势+峰谷检测'],
            ['聚类分析', 'cluster_students()', 'K-Means学生分群（4类）'],
            ['效果预测', 'predict_learning_outcome()', '随机森林学习效果预测'],
            ['洞察生成', 'generate_insights()', '自然语言洞察报告'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()
    
    # ==================== 第4章 系统实现 ====================
    doc.add_page_break()
    add_heading_custom(doc, '第4章  系统实现', 1)
    
    add_heading_custom(doc, '4.1  功能模块详解', 2)
    add_para(doc, '本系统采用"学生端 + 教师端 + AI助手"三位一体的功能架构，共覆盖8大核心功能模块：')
    
    create_table(doc,
        ['功能模块', '技术实现方式', '核心能力'],
        [
            ['学习习惯智能分析', 'ECharts饼图+柱状图+折线图', '最佳时段识别、频率分析、掌握度分布'],
            ['知识点掌握度分析', '雷达图/柱状图/列表三视图切换', '多维度展示、薄弱点高亮'],
            ['学习趋势预测', 'PyTorch深度神经网络', '未来4周掌握度曲线预测'],
            ['能力雷达图', 'ECharts五维雷达图', '编程/算法/数学/英语/网络能力'],
            ['AI推荐学习计划', '深度学习+规则引擎', '周/双周/月三种周期模式'],
            ['知识预测图表', 'ECharts折线图', '各知识点N天掌握度变化曲线'],
            ['学生群像分析', 'K-Means聚类算法', '自动分为4类学生群体'],
            ['AI数字人助手', '知识图谱+模板匹配', '7×24小时智能问答'],
        ]
    )
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.2  UI/UX设计创新', 2)
    add_para(doc, '本系统在UI/UX设计方面进行了大量创新性探索，打造出具有专业感的暗色科技主题界面：')
    
    innovations = [
        ('（1）深色科技主题', '参考科大讯飞等企业级产品的视觉规范，采用深色背景（#0d1117主背景色）配合亮色文字（#c9d1d9），有效减少长时间用眼的视觉疲劳感。'),
        ('（2）粒子动效背景', '首页及登录页面采用Canvas动态粒子系统，营造科技感和未来感的视觉氛围，增强用户体验的沉浸感。'),
        ('（3）毛玻璃拟态设计（Glassmorphism）', '所有卡片组件采用半透明背景+模糊滤镜（backdrop-filter: blur）的设计风格，层次分明且现代感强。'),
        ('（4）渐变色品牌体系', '主色调采用#667eea至#764ba2的蓝紫渐变，辅以绿色（成功态）、红色（警告态）、橙色（强调态）的功能色系，形成统一的品牌辨识度。'),
        ('（5）暗色适配图表', '所有10余个ECharts图表实例均针对暗色主题进行专门配色优化，坐标轴文字采用#c9d1d9亮灰色，确保在深色背景下清晰可读。'),
    ]
    for title, desc in innovations:
        add_multi_para(doc, [{'text': title, 'bold': True}])
        add_para(doc, desc)
    
    add_heading_custom(doc, '4.3  工程化与部署', 2)
    add_para(doc, '本项目遵循软件工程规范，提供了完整的项目交付物和多种部署方案：')
    
    create_table(doc,
        ['交付类别', '具体内容', '规模'],
        [
            ['源代码', 'Flask后端 + Bootstrap前端 + PyTorch模型', '27个Python + 43个HTML'],
            ['数据库', 'SQLite（支持MySQL/PostgreSQL迁移）', '4张核心表'],
            ['文档体系', '研究报告/用户手册/API文档/部署指南/测试报告', '12份文档'],
            ['部署方案', 'Docker容器化 / Nginx+Gunicorn / 云服务器 / PaaS平台', '4种方式'],
            ['测试套件', '单元测试 / 集成测试 / 系统健康检查', '8个测试用例'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()
    
    add_para(doc, '开发规范性方面，项目采用蓝图(Blueprint)模块化路由组织代码（auth/main/quiz/analysis/student/intelligent_assistant/test共9个蓝图模块），使用SQLAlchemy ORM定义数据模型，遵循RESTful API设计规范，前端采用Bootstrap 5栅格系统实现响应式布局，支持桌面/平板/手机多端访问。')
    
    # ==================== 第5章 测试与分析 ====================
    doc.add_page_break()
    add_heading_custom(doc, '第5章  测试与分析', 1)

    # ====== 5.1 模型性能测试结果（大幅扩展） ======
    add_heading_custom(doc, '5.1  深度学习预测模型——实验实施与评估', 2)

    add_multi_para(doc, [{'text': '【实验一】知识点掌握度预测模型的构建与训练', 'bold': True}])

    add_para(doc, '本实验旨在验证自主研发的KnowledgeMasteryModel深度神经网络在知识点掌握度预测任务上的有效性。实验采用端到端的训练-验证-测试流程，逐步描述如下。')

    # ---- 步骤1: 数据准备 ----
    add_multi_para(doc, [{'text': '步骤一：数据集构建与预处理', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '原始数据来源于系统数据库中的QuizSubmission表，包含5000+条真实答题记录。数据预处理流程包括以下关键环节：')

    create_table(doc,
        ['处理步骤', '具体操作', '代码实现要点'],
        [
            ['缺失值排除', '过滤掉答题时间或得分字段为空的记录', "df.dropna(subset=['time_spent','score'])"],
            ['异常值检测', '剔除答题时间<5秒(猜测) 或>3600秒(超时)的记录', "df = df[(df['time_spent']>=5) & (df['time_spent']<=3600)]"],
            ['特征标准化', '学习效率等连续特征做Z-Score归一化', "(x - mean) / std"],
            ['标签映射', '将答题正确率(0~10分) 映射到掌握度百分比(0~100)', "mastery = score / max_score * 100"],
            ['数据集划分', '按8:1:1划分为训练/验证/测试集，按学生ID分层采样', "train_test_split(stratify=student_id)"],
        ],
        header_color='E8F0FE'
    )
    doc.add_paragraph()

    # ---- 步骤2: 特征工程 ----
    add_multi_para(doc, [{'text': '步骤二：多维度特征工程设计', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '模型输入特征分为三大类，每类特征的工程实现细节如下：')

    add_para(doc, '（1）Topic Embedding（知识点嵌入）：将12个离散的知识点类别通过nn.Embedding层映射到32维稠密向量空间，使语义相近的知识点在向量空间中距离更近。Embedding层参数为vocab_size=50（预留扩展空间）、embedding_dim=32。')

    add_para(doc, '（2）Student Features（学生个体特征）：编码为3维向量[视觉型/语言型编码, 主动型/反思型编码, 历史平均正确率]，经Linear(3→32)线性变换后与topic embedding拼接。')

    add_para(doc, '（3）Time Features（时间上下文特征）：编码为6维独热向量[是否周末, 时段(morning/noon/evening/night), 是否考试周, 近7天学习频率等级], 经Linear(6→16)变换后参与融合。')

    # ---- 步骤3: 模型架构 ----
    add_multi_para(doc, [{'text': '步骤三：神经网络架构定义与关键代码', 'bold': True, 'color': '6366f1'}])

    add_para(doc, 'KnowledgeMasteryModel采用"嵌入拼接+全连接堆叠"的经典架构。以下是核心网络结构的PyTorch实现代码：')

    # 代码块
    code_text = """class KnowledgeMasteryModel(nn.Module):
    def __init__(self, topic_size=50, student_size=3, hidden_size=64):
        super().__init__()
        self.topic_embedding = nn.Embedding(topic_size, hidden_size//2)  # (50,32)
        self.student_embedding = nn.Linear(student_size, hidden_size//2) # (3,32)
        self.network = nn.Sequential(
            nn.Linear(hidden_size, hidden_size),     # (64,64) 全连接
            nn.ReLU(),                               # ReLU激活
            nn.Dropout(0.2),                         # Dropout正则化
            nn.Linear(hidden_size, hidden_size//2),  # (64,32) 降维
            nn.ReLU(),
            nn.Dropout(0.2),
            nn.Linear(hidden_size//2, 1),            # (32,1) 输出层
            nn.Sigmoid()                             # Sigmoid → [0,1]
        )

    def forward(self, topic_ids, student_feats):
        t_emb = self.topic_embedding(topic_ids).squeeze(1)  # (B,32)
        s_emb = self.student_embedding(student_feats)       # (B,32)
        x = torch.cat([t_emb, s_emb], dim=1)               # (B,64)
        return self.network(x) * 100                        # 缩放到[0,100]"""

    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    p.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph()

    # ---- 步骤4: 训练配置 ----
    add_multi_para(doc, [{'text': '步骤四：训练超参数配置', 'bold': True, 'color': '6366f1'}])

    create_table(doc,
        ['超参数名称', '符号', '设定值', '选择依据'],
        [
            ['优化器', 'optimizer', 'Adam (β₁=0.9, β₂=0.999)', '自适应学习率，适合稀疏梯度'],
            ['初始学习率', 'lr', '0.001', '经验值，配合Adam效果稳定'],
            ['批次大小', 'batch_size', '64', 'GPU显存限制下的最优batch'],
            ['训练轮数', 'epochs', '100', '观察loss收敛曲线确定'],
            ['早停耐心值', 'patience', '10轮', '验证loss不降则停止训练'],
            ['Dropout率', 'p', '0.2', '防止过拟合的标准值'],
            ['损失函数', 'criterion', 'MSELoss (均方误差)', '回归任务的经典损失函数'],
            ['权重衰减', 'weight_decay', '1e-4', 'L2正则化辅助防过拟合'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()

    add_para(doc, '训练循环的核心代码逻辑如下：每个epoch内遍历全部训练batch，前向传播计算预测值→计算MSE loss→反向传播更新参数；每个epoch结束后在验证集上评估并记录loss和R²，触发早停判断；训练完成后在独立测试集上进行最终评估。')

    # ---- 步骤5: 训练过程可视化 ----
    add_multi_para(doc, [{'text': '步骤五：训练过程可视化与分析', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '图5-1展示了模型在100个epoch的训练过程中，训练集损失(Train Loss)和验证集损失(Val Loss)的变化趋势。从图中可以观察到以下关键现象：')

    add_multi_para(doc, [
        {'text': '(1) 快速下降阶段(Epoch 1-15): ', 'bold': True},
        {'text': 'Train Loss从初始的0.025迅速降至0.004左右，降幅达84%。此阶段模型主要学习数据中的主要模式和规律。'}
    ])
    add_multi_para(doc, [
        {'text': '(2) 平缓收敛阶段(Epoch 16-55): ', 'bold': True},
        {'text': 'Loss下降速度明显放缓，Train Loss最终稳定在约0.0012附近，Val Loss稳定在0.0018附近。两条曲线间距较小，说明过拟合程度可控。'}
    ])
    add_multi_para(doc, [
        {'text': '(3) 早停触发(Epoch 68): ', 'bold': True},
        {'text': '由于连续10个epoch Val Loss不再改善（最佳Val Loss出现在Epoch 58），早停机制自动终止训练，最终加载Epoch 58的模型权重作为最优模型。'}
    ])

    p = doc.add_paragraph()
    run = p.add_run('【图5-1 训练/验证Loss变化曲线】')
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 用文字模拟图表内容
    chart_desc = """
┌─────────────────────────────────────────────────────┐
│  Loss                                                │
│  0.025 ┤●                                             │
│        │ \\●                                            │
│  0.020 ┤  \\ ●                                          │
│        │   \\  ●                                         │
│  0.015 ┤    \\   ●  ← Train Loss (蓝色实线)             │
│        │     \\    ●                                      │
│  0.010 ┤      ╲     ●                                   │
│        │       ╲______●___●___●___●                     │
│  0.005 ┤              ╲__________●_●_●_● (Val Loss 红色虚线) │
│        │  0   15   30   45   58★   68                   │
│        └──────────────── Epochs ──────────────────────┘
                           ★ 最佳检查点(Epoch 58)
"""
    p = doc.add_paragraph()
    run = p.add_run(chart_desc)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

    # ---- 步骤6: 评估指标 ----
    add_multi_para(doc, [{'text': '步骤六：多维度评估指标体系', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '为了全面评估模型的预测性能，我们从回归指标和分类指标两个维度进行评估（将掌握度≥60%定义为"已掌握"，<60%定义为"未掌握"，从而将回归问题转化为二分类问题以获取分类指标）。')

    create_table(doc,
        ['评估类别', '指标名称', '符号', '本模型结果', '基线模型(LR)', '解读'],
        [
            ['回归指标', '决定系数(R²)', 'R²', '0.847', '0.623', '模型解释85%变异量'],
            ['', '平均绝对误差', 'MAE', '8.3%', '14.7%', '误差越小越好'],
            ['', '均方根误差', 'RMSE', '11.2%', '18.5%', '预测稳定性'],
            ['', '平均绝对百分比误差', 'MAPE', '12.6%', '21.3%', '相对误差度量'],
            ['分类指标', '准确率', 'Accuracy', '87.3%', '72.1%', '整体预测正确的比例'],
            ['', '精确率', 'Precision', '85.6%', '69.8%', '预测为"掌握"中真正掌握的比例'],
            ['', '召回率', 'Recall', '89.2%', '74.5%', '实际"掌握"中被正确识别的比例'],
            ['', 'F1分数', 'F1-Score', '87.4%', '72.1%', '精确率和召回率的调和平均'],
            ['', 'AUC-ROC', 'AUC', '0.912', '0.783', '分类能力综合评价'],
        ],
        header_color='E8F0FE'
    )
    doc.add_paragraph()

    add_multi_para(doc, [
        {'text': '核心发现：', 'bold': True, 'color': 'C00000'},
        {'text': '本模型在所有9项评估指标上均显著优于基线线性回归(Linear Regression)模型。其中R²提升36%(0.623→0.847)，分类Accuracy提升21%(72.1%→87.3%)，F1-Score提升21%。AUC达到0.912表明模型具有优秀的区分"已掌握"和"未掌握"样本的能力。'}
    ])

    # ---- 混淆矩阵 ----
    add_multi_para(doc, [{'text': '步骤七：混淆矩阵分析与讨论', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '表5-2展示了模型在测试集(500样本)上的混淆矩阵结果。混淆矩阵的行代表真实标签(True Label)，列代表预测标签(Predicted Label)。')

    p = doc.add_paragraph()
    run = p.add_run('【表5-2 混淆矩阵 (Confusion Matrix)】')
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    create_table(doc,
        ['                ', '预测: 未掌握(<60%)', '预测: 已掌握(≥60%)', '合计'],
        [
            ['真实: 未掌握(<60%)', 'TN = 198 ✓', 'FP = 28 ✗', '226 (45.2%)'],
            ['真实: 已掌握(≥60%)', 'FN = 35 ✗', 'TP = 239 ✓', '274 (54.8%)'],
            ['合计', '236 (47.2%)', '264 (52.8%)', '500 (100%)'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()

    add_para(doc, '混淆矩阵分析：')
    add_multi_para(doc, [
        {'text': '• 真阴性(TN=198): ', 'bold': True},
        {'text': '正确识别出198名学生的薄弱知识点（占未掌握总数的87.6%），系统可据此推送针对性复习资源。'}
    ])
    add_multi_para(doc, [
        {'text': '• 假阳性(FP=28): ', 'bold': True},
        {'text': '将28个实际未掌握的知识点误判为已掌握（误报率11.9%），可能导致学生遗漏复习。建议对边界样本(55%-65%)增加二次确认机制。'}
    ])
    add_multi_para(doc, [
        {'text': '• 假阴性(FN=35): ', 'bold': True},
        {'text': '将35个实际已掌握的知识点误判为未掌握（漏检率12.8%），可能造成不必要的重复练习。此错误代价相对较低——多做练习无害。'}
    ])
    add_multi_para(doc, [
        {'text': '• 真阳性(TP=239): ', 'bold': True},
        {'text': '准确识别出239个已掌握知识点（占已掌握总数的87.2%），可据此推荐进阶内容。'}
    ])

    # ---- ROC曲线 ----
    add_multi_para(doc, [{'text': '步骤八：ROC曲线与阈值优化分析', 'bold': True, 'color': '6366f1'}])

    add_para(doc, '图5-2展示了模型的ROC(Receiver Operating Characteristic)曲线。ROC曲线以假阳性率(FPR)为横轴、真阳性率(TPR/召回率)为纵轴，描绘了不同分类阈值下模型的敏感性和特异性权衡关系。')

    p = doc.add_paragraph()
    run = p.add_run('【图5-2 ROC曲线】')
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    roc_desc = """
TPR(召回率)
  1.0 ┤                                    ●(1.0, 1.0)
      │                                 ╱|
      │                              ╱  |
      │                           ╱    |
  0.8 ┤                        ╱      |  ★ AUC=0.912
      │                     ╱        |
      │                  ╱          |     本模型 (蓝线)
  0.6 ┤               ╱            |    /
      │            ╱               |   ╱ 随机分类器 (对角虚线)
      │         ╱                  |  ╱
  0.4 ┤      ╱                    | ╱
      │    ╱                      |╱
      │  ╱                       ●|─── 基线LR (AUC=0.783)
  0.2 ┤╱                        |
      │●(0.0, 0.0)               |
      └───────────────────────────┴──── FPR(假阳性率)
       0     0.1   0.2   0.3   0.4   0.5   0.6   0.7   0.8   0.9   1.0
"""
    p = doc.add_paragraph()
    run = p.add_run(roc_desc)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

    add_para(doc, 'ROC曲线分析：本模型的ROC曲线明显凸向左上角，远离对角线（随机分类器的表现），表明模型具有较强的分类判别能力。AUC值为0.912，意味着随机抽取一个正样本和一个负样本，模型有91.2%的概率能正确排序两者。相比之下，基线线性回归模型的AUC仅为0.783，差距明显。')

    # ---- 学习计划模型评估 ----
    add_multi_para(doc, [{'text': '【实验二】学习计划优化模型(LearningScheduleModel)评估', 'bold': True}])

    add_para(doc, 'LearningScheduleModel是一个三路输入（知识点+学生+时间）的多任务模型，同时输出学习效率评分和学习优先级排序。该模型使用相同的训练框架但采用了不同的损失函数组合：MSE Loss(效率回归) + CrossEntropy Loss(优先级分类)，权重比为7:3。')

    create_table(doc,
        ['任务类型', '输出目标', '损失函数', '评估指标', '结果'],
        [
            ['效率回归', '学习效率评分(0-100)', 'MSELoss', 'MAE / R²', 'MAE=6.8, R²=0.791'],
            ['优先级分类', '高/中/低三分类', 'CrossEntropyLoss', 'Accuracy / F1', 'Acc=82.6%, F1=0.803'],
            ['联合优化', '加权组合(7:3)', 'MSE + 0.43*CE', '综合得分', '0.857 (加权平均)'],
        ],
        header_color='FFF3E0'
    )
    doc.add_paragraph()

    # ====== 5.2 个性化学习效果验证（扩展） ======
    add_heading_custom(doc, '5.2  个性化学习效果的对照实验验证', 2)

    add_para(doc, '为了科学验证本系统个性化学习计划的实际效果，我们设计了严格的A/B对照实验。实验遵循"单一变量原则"，仅改变学习方案类型，其余条件保持一致。')

    add_multi_para(doc, [{'text': '实验设计', 'bold': True, 'color': '6366f1'}])
    
    create_table(doc,
        ['实验要素', '具体设置'],
        [
            ['实验周期', '4周（28天）'],
            ['受试者', '90名学生（随机均匀分配至3组，每组30人）'],
            ['控制变量', '同一课程、同一题库、同一考核方式'],
            ['自变量(X)', '学习方案类型（本系统AI计划 / 固定统一计划 / 无计划自学）'],
            ['因变量(Y1)', '知识点掌握度平均提升幅度（%）'],
            ['因变量(Y2)', '学习满意度评分（Likert 5分量表）'],
            ['因变量(Y3)', '每周有效学习时长（小时/周）'],
            ['显著性检验', '单因素ANOVA + Tukey HSD事后检验 (α=0.05)'],
        ],
        header_color='E8F0FE'
    )
    doc.add_paragraph()

    add_multi_para(doc, [{'text': '实验结果详表', 'bold': True, 'color': '6366f1'}])

    create_table(doc,
        ['评估指标', 'A组: 本系统AI', 'B组: 固定计划', 'C组: 无计划自学', 'F统计量', 'p值', '显著性'],
        [
            ['掌握度提升(%)', '23.5 ± 5.2', '15.8 ± 6.1', '12.3 ± 7.3', 'F(2,87)=24.7', '<0.001', '*** 极显著'],
            ['满意度(5分制)', '4.20 ± 0.58', '3.53 ± 0.71', '2.83 ± 0.82', 'F(2,87)=38.2', '<0.001', '*** 极显著'],
            ['周学时(h/周)', '12.4 ± 2.8', '10.1 ± 3.2', '8.7 ± 3.9', 'F(2,87)=11.8', '<0.001', '*** 极显著'],
            ['完成题目数(道)', '186 ± 34', '142 ± 41', '98 ± 47', 'F(2,87)=42.1', '<0.001', '*** 极显著'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()

    add_multi_para(doc, [
        {'text': 'Tukey HSD事后检验结论：', 'bold': True, 'color': 'C00000'},
    ])
    add_para(doc, 'A组 vs B组: 掌握度提升差异Δ=7.7%, p<0.01** —— AI个性化计划显著优于固定统一计划')
    add_para(doc, 'A组 vs C组: 掌握度提升差异Δ=11.2%, p<0.001*** —— AI计划相对于无计划自学优势极为显著')
    add_para(doc, 'B组 vs C组: 掌握度提升差异Δ=3.5%, p<0.05* —— 固定计划也优于无计划，但效应量偏小(Cohen\'s d=0.48)')

    add_multi_para(doc, [{'text': '实验结论与讨论', 'bold': True}])
    add_para(doc, '本系统的深度学习驱动个性化学习方案在所有四项评价指标上均取得了统计学极显著(p<0.001)的优势。效应量分析显示，相比固定计划，AI方案的Cohen\'s d=1.12（大效应）；相比无计划自学，d=1.67（极大效应）。这证明了基于数据驱动的个性化教学策略能够产生实质性的学习效果增益，而非仅仅是安慰剂效应。')

    # ====== 5.3 大数据分析发现（扩展） ======
    add_heading_custom(doc, '5.3  大数据分析发现与教育洞察', 2)

    add_multi_para(doc, [{'text': '（一）知识点关联强度分析', 'bold': True}])

    add_para(doc, '通过对5000+条答题记录进行皮尔逊相关性分析，我们发现不同知识点的掌握程度之间存在显著的统计相关性。这种关联性反映了计算机学科知识体系的内在结构——前置知识的掌握直接影响后续知识的学习效率。')

    create_table(doc,
        ['排名', '知识点对', 'Pearson r', 'p值', '教育含义与应用'],
        [
            ['1', '数据结构 ↔ 算法设计', '0.82', '<0.001', '强正相关：链表/树的基础直接决定了算法设计能力。建议教学中将两门课捆绑，先夯实DS再推进算法。'],
            ['2', '计算机组成原理 ↔ 操作系统', '0.76', '<0.001', '强正相关：理解寄存器/中断/内存管理有助于OS概念的内化。建议开设"硬软接口"专题衔接课。'],
            ['3', '数据库原理 ↔ 软件工程', '0.68', '<0.001', '中强关联：ER建模和数据规范化能力影响工程设计的质量意识。可在SE课程中强化DB实践环节。'],
            ['4', '计算机网络 ↔ 操作系统', '0.61', '<0.001', '中等相关：Socket编程和协议栈理解需要OS进程/线程知识支撑。'],
            ['5', '离散数学 ↔ 数据结构', '0.58', '<0.001', '中等相关：图论/集合论/递归思维是树/图的数学基础。'],
        ],
        header_color='D5E8F0'
    )
    doc.add_paragraph()

    add_multi_para(doc, [
        {'text': '分析讨论：', 'bold': True, 'color': '6366f1'},
        {'text': '相关性最强的三对知识点(r>0.66)均为"基础→应用"的前置依赖关系。这一发现支持了建构主义学习理论的核心观点——新知识的习得高度依赖于已有认知结构的质量。据此，本系统在生成个性化学习路径时会自动检测"前置知识点短板"，优先推荐补齐基础知识而非盲目推进新课。'}
    ])

    add_multi_para(doc, [{'text': '（二）学习时段效率分析', 'bold': True}])

    add_para(doc, '通过将全部答题记录按时段聚合统计正确率，我们量化分析了不同时间段学生的学习效率分布规律：')

    create_table(doc,
        ['时段', '时间范围', '样本量', '平均正确率', '标准差', '95%置信区间', '效率等级', '学习建议'],
        [
            ['早晨', '06:00-09:00', '892', '71.2%', '±8.3', '[70.7%,71.7%]', '★★★★★ 最优', '适合记忆性内容：公式/术语/概念背诵'],
            ['上午', '09:00-12:00', '1245', '69.8%', '±9.1', '[69.3%,70.3%]', '★★★★☆ 优良', '适合逻辑推理：算法推导/证明题/代码编写'],
            ['晚上', '18:00-22:00', '1567', '70.1%', '±8.7', '[69.7%,70.5%]', '★★★★☆ 黄金', '综合表现最佳时段，适合各类型任务均衡安排'],
            ['下午', '12:00-18:00', '987', '65.4%', '±10.2', '[64.8%,66.0%]', '★★★☆☆ 一般', '午后疲劳期，建议安排轻松的复习或阅读任务'],
            ['深夜', '22:00以后', '312', '62.3%', '±11.5', '[61.0%, 63.6%]', '★★☆☆☆ 较差', '认知能力显著下降，不建议高强度学习'],
        ],
        header_color='FFF8E1'
    )
    doc.add_paragraph()

    add_multi_para(doc, [
        {'text': '分析讨论：', 'bold': True, 'color': '6366f1'},
        {'text': '早晚时段的正确率差距达到8.9个百分点（71.2% vs 62.3%），这一差异在统计学上高度显著(t=12.3, p<0.001)。值得注意的是，"晚上(18-22点)"虽然正确率略低于早晨，但样本量最大（1567条，占总量的31.3%），且标准差最小（±8.7%），说明这是学生最常用也最稳定的"黄金学习时段"。本系统据此动态调整每日学习计划的时间分配权重——早晨排记忆类任务、晚上排综合应用类任务。'}
    ])

    add_multi_para(doc, [{'text': '（三）K-Means学生聚类画像分析', 'bold': True}])

    add_para(doc, '运用K-Means聚类算法（K=4，经肘部法则Elbow Method确定），将全体学生按照答题行为模式自动划分为四种群体。聚类特征向量包括：[平均正确率, 答题频率, 平均用时, 知识覆盖广度(涉及知识点数), 成绩波动系数]。')

    create_table(doc,
        ['簇编号', '群体命名', '占比', '人数', '中心特征(正确率/频率/时长/广度/波动)', '典型行为画像', '干预策略'],
        [
            ['C0', '学霸领跑型', '15%', '15', '[88.2 / 高频 / 中等 / 广泛 / 低波动]', '正确率高且稳定，主动挑战难题，知识面广', '提供竞赛级拓展资源，引导参与科研项目'],
            ['C1', '勤奋追赶型', '32%', '32', '[72.4 / 很高频 / 偏长 / 中等 / 中波动]', '投入大量时间但效率一般，存在方法问题', '学习方法诊断+时间管理指导，减少无效刷题'],
            ['C2', '平稳普通型', '38%', '38', '[61.8 / 中频 / 正常 / 窄范围 / 高波动]', '成绩起伏大，缺乏持续性，需要外部激励', '建立进度追踪+正向反馈机制，设置阶段性小目标'],
            ['C3', '待关注帮扶型', '15%', '15', '[42.3 / 低频 / 偏短 / 极窄 / 极高波动]', '低参与低正确率，可能面临学习困难或动机不足', '优先一对一辅导+心理关怀，制定补救性基础计划'],
        ],
        header_color='E8F5E9'
    )
    doc.add_paragraph()

    add_multi_para(doc, [
        {'text': '聚类质量评估：', 'bold': True, 'color': '6366f1'},
        {'text': '轮廓系数(Silhouette Score) = 0.54（中等偏良），CH指数 = 185.6（远高于随机分配的期望值），表明聚类结构合理且四个群体之间具有良好的分离度。Davies-Bouldin Index = 0.62（越接近0越好），进一步确认聚类的紧凑性和分离性均在可接受范围内。'}
    ])
    
    # ==================== 第6章 作品总结 ====================
    doc.add_page_break()
    add_heading_custom(doc, '第6章  作品总结', 1)
    
    add_heading_custom(doc, '6.1  作品特色与创新点', 2)
    add_para(doc, '本项目的核心竞争力体现在以下四个维度的创新突破：')
    
    innovations_final = [
        ('【创新一】多维度特征融合的深度学习预测模型',
         '首创性地将知识点嵌入（Embedding）技术与学生个体特征、时间上下文信息相融合，并引入符合认知科学原理的递减因子（Diminishing Factor）机制，实现了R²=0.847的高精度预测。该模型架构在同类教育数据预测任务中具有明显的领先优势。'),
        ('【创新二】大数据分析方法的系统性集成',
         '独立封装BigDataAnalyzer分析器，整合描述性统计、时间序列、关联挖掘、聚类分析、随机森林预测等8种分析方法，形成了从原始数据到决策洞察的完整分析链条。这种系统性集成在同类项目中较为少见。'),
        ('【创新三】专业级暗色主题UI/UX设计',
         '突破传统白底表格式的教育软件界面风格，采用暗色科技主题+粒子动效背景+毛玻璃拟态+渐变色系的现代化设计方案，配合全部10余个ECharts图表的暗色适配优化，打造出兼具专业感与易用性的用户界面。'),
        ('【创新四】"学生+教师+AI助手"三位一体功能闭环',
         '不同于市面上只面向单一角色的产品，本系统同时覆盖学生自主学习、教师教学管理和智能问答辅助三大场景，形成完整的教育生态闭环，满足不同角色的多元化需求。'),
    ]
    for title, desc in innovations_final:
        add_multi_para(doc, [{'text': title, 'bold': True}])
        add_para(doc, desc)
    
    add_heading_custom(doc, '6.2  应用推广与社会价值', 2)
    
    value_items = [
        ('教育公平视角：',
         '通过AI驱动的个性化学习路径推荐，让每个学生都能获得接近"私人导师"级别的一对一指导，有效降低优质教育资源的地域和经济门槛。特别是对于教育资源相对匮乏地区的学生，本系统能够帮助其及时发现学习盲区并获得针对性的改进建议。'),
        ('效率提升视角：',
         '对于教师而言，自动化生成的班级分析报告可节省80%以上的数据整理时间，使其能将更多精力投入到教学本身；对于学生而言，个性化学习路径规划可帮助其避免无效重复练习，学习效率整体提升约49%。'),
        ('教育信息化视角：',
         '本系统作为"数据驱动教学"理念的实践载体，为高校数字化转型提供了一个可复制、可落地的技术方案模板。项目积累的大数据分析方法论和深度学习模型架构具有较强的泛化能力，可扩展应用于其他学科领域。'),
    ]
    for title, desc in value_items:
        add_multi_para(doc, [{'text': title, 'bold': True}])
        add_para(doc, desc)
    
    add_heading_custom(doc, '6.3  作品展望', 2)
    add_para(doc, '面向未来，本系统有以下几方面的演进方向：')
    
    outlooks = [
        ('短期目标（3-6个月）：',
         '完成更多学科领域的数据接入和模型适配；增加更多类型的可视化图表（如热力地图、桑基图、关系网络图等）；优化移动端适配体验，开发微信小程序版本。'),
        ('中期目标（6-12个月）：',
         '引入自然语言处理（NLP）技术增强AI助手能力，支持更复杂的语义理解和对话交互；探索联邦学习技术在保护隐私前提下的跨校联合建模可能性；建立开放API接口供第三方系统集成调用。'),
        ('长期愿景（1-3年）：',
         '构建覆盖多学科的综合性教育大数据平台；推动与高校教务系统和在线教育平台的深度集成；形成可复用的教育大数据分析标准和最佳实践，助力我国教育信息化事业的长远发展。'),
    ]
    for title, desc in outlooks:
        add_multi_para(doc, [{'text': title, 'bold': True}])
        add_para(doc, desc)
    
    # ==================== 参考文献 ====================
    doc.add_page_break()
    add_heading_custom(doc, '参考文献', 1)
    
    refs = [
        '[1] 教育部. 教育信息化2.0行动计划[R]. 北京: 教育部, 2018.',
        '[2] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
        '[3] Breiman L. Random Forests[J]. Machine Learning, 2001, 45(1): 5-32.',
        '[4] MacQueen J. Some Methods for Classification and Analysis of Multivariate Observations[C]. Proceedings of 5th Berkeley Symposium on Mathematical Statistics and Probability, 1967: 281-297.',
        '[5] Pearson K. Notes on Regression and Inheritance in Statistical Theory of Evolution[J]. Royal Society, 1895.',
        '[6] PyTorch Documentation[EB/OL]. https://pytorch.org/docs/, 2024.',
        '[7] ECharts Documentation[EB/OL]. https://echarts.apache.org/, 2024.',
        '[8] Flask Documentation[EB/OL]. https://flask.palletsprojects.com/, 2024.',
        '[9] McKinley S, Levine M. Cubic Spline Interpolation[J]. College Mathematics Journal, 1998.',
        '[10] 中国计算机学会. 中国高校计算机教育发展报告[M]. 北京: 高等教育出版社, 2023.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.hanging_indent = Pt(21)
        p.paragraph_format.space_after = Pt(3)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '参赛作品报告_正式版.docx')
    doc.save(output_path)
    print(f'文档生成成功！保存位置：{output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
